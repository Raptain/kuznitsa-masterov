from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import string

class DocGenerator:
    def __init__(self, output_dir='static/docs'):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate_doc(self, rooms, connections, map_image_path, filename='dungeon_report.docx'):
        """Создаёт подробный отчёт с картой, комнатами и связями"""
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('КУЗНИЦА МАСТЕРОВ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading('Отчёт о подземелье', 1)
        
        # Добавляем карту
        doc.add_heading('Карта подземелья', level=2)
        doc.add_picture(map_image_path, width=Inches(6))
        doc.add_paragraph()
        
        # ============= СТАТИСТИКА =============
        doc.add_heading('Статистика', level=2)
        
        # Подсчёт типов комнат
        room_types_count = {}
        for room in rooms:
            room_type = room['type']
            room_types_count[room_type] = room_types_count.get(room_type, 0) + 1
        
        # Названия типов для отображения
        type_names = {
            'entrance': 'Вход', 'boss': 'Босс', 'battle': 'Битва',
            'treasure': 'Сокровище', 'trap': 'Ловушка', 'puzzle': 'Головоломка',
            'rest': 'Отдых', 'gateway': 'Портал', 'shop': 'Кузница'
        }
        
        # Таблица статистики
        stats_table = doc.add_table(rows=len(room_types_count) + 3, cols=2)
        stats_table.style = 'Table Grid'
        
        # Заголовки
        stats_table.cell(0, 0).text = 'Показатель'
        stats_table.cell(0, 1).text = 'Значение'
        
        # Основная статистика
        stats_table.cell(1, 0).text = 'Всего комнат'
        stats_table.cell(1, 1).text = str(len(rooms))
        
        stats_table.cell(2, 0).text = 'Всего связей'
        stats_table.cell(2, 1).text = str(len(connections))
        
        row = 3
        for room_type, count in room_types_count.items():
            type_name = type_names.get(room_type, room_type)
            stats_table.cell(row, 0).text = f'Комнат типа "{type_name}"'
            stats_table.cell(row, 1).text = str(count)
            row += 1
        
        doc.add_paragraph()
        
        # ============= ТАБЛИЦА КОМНАТ =============
        doc.add_heading('Список комнат', level=2)
        
        room_table = doc.add_table(rows=1, cols=5)
        room_table.style = 'Table Grid'
        room_table.autofit = False
        
        # Устанавливаем ширину колонок
        room_table.columns[0].width = Inches(0.5)   # №
        room_table.columns[1].width = Inches(1.0)   # Тип
        room_table.columns[2].width = Inches(0.8)   # Размер
        room_table.columns[3].width = Inches(0.8)   # Сложность
        room_table.columns[4].width = Inches(3.5)   # Описание
        
        # Заголовки таблицы комнат
        hdr = room_table.rows[0].cells
        hdr[0].text = '№'
        hdr[1].text = 'Тип'
        hdr[2].text = 'Размер'
        hdr[3].text = 'Сложность'
        hdr[4].text = 'Описание'
        
        # Заполняем комнаты
        for room in rooms:
            row_cells = room_table.add_row().cells
            row_cells[0].text = str(room['id'])
            row_cells[1].text = type_names.get(room['type'], room['type'])
            row_cells[2].text = room.get('size', 'medium')
            row_cells[3].text = str(room.get('difficulty', 1))
            row_cells[4].text = room.get('description', 'Нет описания')
        
        doc.add_paragraph()
        
        # ============= ТАБЛИЦА СВЯЗЕЙ =============
        doc.add_heading('Связи между комнатами', level=2)
        
        # Создаём таблицу связей
        conn_table = doc.add_table(rows=1, cols=6)
        conn_table.style = 'Table Grid'
        conn_table.autofit = False
        
        # Устанавливаем ширину колонок
        conn_table.columns[0].width = Inches(0.6)   # Буква
        conn_table.columns[1].width = Inches(0.8)   # Из комнаты
        conn_table.columns[2].width = Inches(0.8)   # В комнату
        conn_table.columns[3].width = Inches(1.2)   # Тип связи
        conn_table.columns[4].width = Inches(1.5)   # Направление
        conn_table.columns[5].width = Inches(2.5)   # Описание
        
        # Заголовки таблицы связей
        conn_hdr = conn_table.rows[0].cells
        conn_hdr[0].text = 'Обозначение'
        conn_hdr[1].text = 'Из комнаты'
        conn_hdr[2].text = 'В комнату'
        conn_hdr[3].text = 'Тип связи'
        conn_hdr[4].text = 'Направление'
        conn_hdr[5].text = 'Описание прохода'
        
        # Генерируем буквы для связей (A, B, C...)
        tunnel_letters = {}
        letter_index = 0
        
        # Типы связей с описаниями
        type_names_conn = {
            'normal': 'Обычный проход',
            'secret': 'Секретный проход',
            'oneway': 'Односторонний проход',
            'magic': 'Магический портал'
        }
        
        type_directions = {
            'normal': 'Двусторонний',
            'secret': 'Двусторонний (скрытый)',
            'oneway': 'Только вперёд',
            'magic': 'Мгновенный (туда/обратно)'
        }
        
        for conn in connections:
            # Присваиваем букву для каждой связи
            conn_key = f"{conn['from']}-{conn['to']}"
            if conn_key not in tunnel_letters:
                tunnel_letters[conn_key] = string.ascii_uppercase[letter_index % 26]
                letter_index += 1
            letter = tunnel_letters[conn_key]
            
            row_cells = conn_table.add_row().cells
            row_cells[0].text = letter
            row_cells[1].text = str(conn['from'])
            row_cells[2].text = str(conn['to'])
            row_cells[3].text = type_names_conn.get(conn['type'], conn['type'])
            row_cells[4].text = type_directions.get(conn['type'], 'Неизвестно')
            row_cells[5].text = conn.get('label', 'Обычный коридор')
        
        doc.add_paragraph()
        
        # ============= ТЕКСТОВАЯ СХЕМА СВЯЗЕЙ =============
        doc.add_heading('Схема связей', level=2)
        
        # Группируем связи по исходной комнате
        from collections import defaultdict
        connections_by_from = defaultdict(list)
        for conn in connections:
            connections_by_from[conn['from']].append(conn)
        
        # Создаём текстовую схему
        for from_room in sorted(connections_by_from.keys()):
            conns = connections_by_from[from_room]
            # Находим буквы для связей
            letters = []
            for conn in conns:
                conn_key = f"{conn['from']}-{conn['to']}"
                letter = tunnel_letters.get(conn_key, '?')
                letters.append(f"{letter} → {conn['to']}")
            
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"Комната {from_room} → ")
            run.bold = True
            paragraph.add_run(", ".join(letters))
        
        doc.add_paragraph()
        
        # ============= СПИСОК СВЯЗЕЙ С ОПИСАНИЯМИ =============
        doc.add_heading('Детальное описание переходов', level=2)
        
        for conn in connections:
            conn_key = f"{conn['from']}-{conn['to']}"
            letter = tunnel_letters.get(conn_key, '?')
            
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"[{letter}] Комната {conn['from']} → Комната {conn['to']}")
            run.bold = True
            paragraph.add_run(f"\n  📍 Тип: {type_names_conn.get(conn['type'], conn['type'])}")
            paragraph.add_run(f"\n  🧭 Направление: {type_directions.get(conn['type'], 'Неизвестно')}")
            paragraph.add_run(f"\n  📝 Описание: {conn.get('label', 'Обычный коридор')}")
            paragraph.add_run(f"\n  {'─' * 50}")
        
        doc.add_paragraph()
        
        # ============= ЛЕГЕНДА =============
        doc.add_heading('Легенда', level=2)
        
        # Легенда типов комнат
        doc.add_paragraph('Типы комнат:', style='List Bullet')
        legend_rooms = [
            '1 - Вход', '2 - Босс', '3 - Битва', '4 - Сокровище',
            '5 - Ловушка', '6 - Головоломка', '7 - Отдых', '8 - Портал', '9 - Кузница'
        ]
        for item in legend_rooms:
            doc.add_paragraph(item, style='List Bullet')
        
        # Легенда типов связей
        doc.add_paragraph('Типы переходов:', style='List Bullet')
        legend_connections = [
            '═══ - Обычный проход (сплошная линия, двусторонний)',
            '- - - - Секретный проход (пунктирная линия, требует поиска)',
            '═══→ - Односторонний проход (со стрелкой, только вперёд)',
            '≈≈≈ - Магический портал (волнистая линия, телепортация)'
        ]
        for item in legend_connections:
            doc.add_paragraph(item, style='List Bullet')
        
        # Информация о буквах
        doc.add_paragraph('Обозначения переходов:', style='List Bullet')
        doc.add_paragraph('A, B, C... - буквенные обозначения конкретных туннелей на схеме', style='List Bullet')
        
        # Сохраняем документ
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        print(f"✅ Отчёт сохранён: {filepath}")
        
        return filepath